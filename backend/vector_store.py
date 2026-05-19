from langchain.text_splitter import TokenTextSplitter
from langchain_community.vectorstores import FAISS
from langchain.docstore.document import Document
from io import BytesIO
from langchain_google_genai import GoogleGenerativeAIEmbeddings
import os
import docx
from PyPDF2 import PdfReader
from dotenv import load_dotenv
from .file import process_document

load_dotenv()

# Configuração dos embeddings
embeddings = GoogleGenerativeAIEmbeddings(model="models/text-multilingual-embedding-002")

def limpar_texto(texto):
    return " ".join(texto.split())

def process_document(files, file_extension):
    raw_text = ""
    if file_extension == "pdf":
        reader = PdfReader(files)
        for page in reader.pages:
            text = page.extract_text()
            if text:
                raw_text += text
    elif file_extension == "txt":
        if isinstance(files, BytesIO):
            files.seek(0)
            raw_text = files.read().decode("utf-8")
        else:
            with open(files, "r") as file:
                raw_text = file.read()
    elif file_extension == "docx":
        doc = docx.Document(files)
        raw_text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
    else:
        raise ValueError("Tipo de arquivo não suportado")

    text_splitter = TokenTextSplitter(chunk_size=350, chunk_overlap=100)
    documents = text_splitter.split_text(raw_text)
    document_objs = [Document(page_content=doc) for doc in documents]
    return document_objs

def criar_indices_faiss(files, file_extension):
    documents = process_document(files, file_extension)
    vector_store = FAISS.from_documents(documents, embeddings)
    vector_store.save_local("faiss_index")
    return vector_store

def adicionar_texto_ao_indice(vector_store, files, file_extension):
    documents = process_document(files, file_extension)
    vector_store.add_documents(documents)
    vector_store.save_local("faiss_index")

def verificar_e_atualizar_indice(files, file_extension):
    if os.path.exists("faiss_index"):
        vector_store = FAISS.load_local("faiss_index", embeddings)
        adicionar_texto_ao_indice(vector_store, files, file_extension)
    else:
        vector_store = criar_indices_faiss(files, file_extension)
    return vector_store

def procurar_similaridade(consulta):
    vector_store = FAISS.load_local("faiss_index", embeddings)
    resultados = vector_store.similarity_search(query=consulta, k=1000)
    textos_resultados = [limpar_texto(doc.page_content) for doc in resultados]
    return "".join(textos_resultados)
