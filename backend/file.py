"""
Extração de texto de diversos formatos de arquivo.
Suporta: PDF, TXT, DOCX, XLSX/XLS, CSV, imagens (passadas direto ao Gemini Vision).
"""
from __future__ import annotations

import csv
import io
from io import BytesIO
from typing import Optional

# ── Tipos suportados ──────────────────────────────────────────────────────────

TEXT_EXTENSIONS = {'pdf', 'txt', 'docx', 'doc', 'xlsx', 'xls', 'csv'}
IMAGE_EXTENSIONS = {'png', 'jpg', 'jpeg', 'gif', 'webp', 'bmp'}
ALL_EXTENSIONS   = TEXT_EXTENSIONS | IMAGE_EXTENSIONS

MIME_TYPES = {
    'png':  'image/png',
    'jpg':  'image/jpeg',
    'jpeg': 'image/jpeg',
    'gif':  'image/gif',
    'webp': 'image/webp',
    'bmp':  'image/bmp',
}

FILE_ICONS = {
    'pdf':  '📄', 'txt': '📝', 'docx': '📝', 'doc': '📝',
    'xlsx': '📊', 'xls': '📊', 'csv':  '📊',
    'png':  '🖼️', 'jpg': '🖼️', 'jpeg': '🖼️',
    'gif':  '🖼️', 'webp':'🖼️', 'bmp':  '🖼️',
}


def get_extension(filename: str) -> str:
    return filename.rsplit('.', 1)[-1].lower() if '.' in filename else ''


def get_mime_type(ext: str) -> str:
    return MIME_TYPES.get(ext.lower(), 'application/octet-stream')


def is_image(ext: str) -> bool:
    return ext.lower() in IMAGE_EXTENSIONS


def extract_text(file_bytes: bytes, ext: str) -> Optional[str]:
    """
    Extrai texto do arquivo e retorna como string.
    Retorna None para imagens (devem ser tratadas com invoke_with_image).
    """
    ext = ext.lower()

    if ext in IMAGE_EXTENSIONS:
        return None  # Usar Gemini Vision

    if ext == 'pdf':
        return _extract_pdf(file_bytes)

    if ext == 'txt':
        return _decode_text(file_bytes)

    if ext in ('docx', 'doc'):
        return _extract_docx(file_bytes)

    if ext in ('xlsx', 'xls'):
        return _extract_excel(file_bytes)

    if ext == 'csv':
        return _extract_csv(file_bytes)

    raise ValueError(f"Tipo de arquivo não suportado: .{ext}")


# ── Extratores internos ───────────────────────────────────────────────────────

def _extract_pdf(data: bytes) -> str:
    from PyPDF2 import PdfReader
    reader = PdfReader(BytesIO(data))
    pages = []
    for i, page in enumerate(reader.pages, 1):
        text = page.extract_text()
        if text and text.strip():
            pages.append(f"[Página {i}]\n{text.strip()}")
    return '\n\n'.join(pages) or '(PDF sem texto extraível)'


def _decode_text(data: bytes) -> str:
    for enc in ('utf-8', 'latin-1', 'cp1252'):
        try:
            return data.decode(enc)
        except UnicodeDecodeError:
            continue
    return data.decode('utf-8', errors='replace')


def _extract_docx(data: bytes) -> str:
    import docx as python_docx
    doc = python_docx.Document(BytesIO(data))
    parts = []

    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)

    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(' | '.join(cells))
        if rows:
            parts.append('\n'.join(rows))

    return '\n'.join(parts) or '(Documento vazio)'


def _extract_excel(data: bytes) -> str:
    import openpyxl
    wb = openpyxl.load_workbook(BytesIO(data), read_only=True, data_only=True)
    parts = []

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows = []
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) if c is not None else '' for c in row]
            if any(c.strip() for c in cells):
                rows.append('\t'.join(cells))
        if rows:
            parts.append(f"=== Planilha: {sheet_name} ===\n" + '\n'.join(rows))

    wb.close()
    return '\n\n'.join(parts) or '(Planilha vazia)'


def _extract_csv(data: bytes) -> str:
    text = _decode_text(data)
    # Detecta delimitador automaticamente
    try:
        dialect = csv.Sniffer().sniff(text[:2048])
    except csv.Error:
        dialect = csv.excel

    reader = csv.reader(io.StringIO(text), dialect)
    rows = ['\t'.join(row) for row in reader if any(c.strip() for c in row)]
    return '\n'.join(rows) or '(CSV vazio)'
